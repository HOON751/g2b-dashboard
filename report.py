"""
보고서 생성 모듈
report.py

검색된 DataFrame을 받아 분석 텍스트와 Word(.docx) 보고서를 생성합니다.
"""

import io
from datetime import datetime


def _won(x):
    try:
        return f"{int(round(float(x))):,}원"
    except Exception:
        return str(x)


def _eok(x):
    try:
        v = float(x)
        if v >= 1e8:
            return f"{v/1e8:,.1f}억원"
        if v >= 1e4:
            return f"{v/1e4:,.0f}만원"
        return f"{v:,.0f}원"
    except Exception:
        return str(x)


def build_analysis(df, term, bgn, end):
    """
    DataFrame을 분석해서 보고서용 구조화 데이터(dict)를 반환.
    화면 표시와 Word 생성에 공통으로 사용.
    """
    a = {}
    total_amt = float(df["공급금액"].sum()) if "공급금액" in df else 0
    n = len(df)
    a["검색어"] = term
    a["기간"] = f"{bgn} ~ {end}"
    a["총공급금액"] = total_amt
    a["계약건수"] = n
    a["업체수"] = int(df["업체명"].nunique()) if "업체명" in df else 0
    a["수요기관수"] = int(df["수요기관"].nunique()) if "수요기관" in df else 0
    a["평균계약금액"] = total_amt / n if n else 0

    # 상위 업체
    if "업체명" in df and "공급금액" in df:
        g = df.groupby("업체명")["공급금액"].agg(["sum", "count"]).sort_values("sum", ascending=False)
        g["점유율"] = g["sum"] / total_amt * 100 if total_amt else 0
        a["상위업체"] = g.head(10).reset_index()

    # 상위 수요기관
    if "수요기관" in df and "공급금액" in df:
        g = df.groupby("수요기관")["공급금액"].agg(["sum", "count"]).sort_values("sum", ascending=False)
        g["점유율"] = g["sum"] / total_amt * 100 if total_amt else 0
        a["상위기관"] = g.head(10).reset_index()

    # 우수제품 비중
    if "우수제품여부" in df and "공급금액" in df:
        tmp = df.copy()
        tmp["_우수"] = tmp["우수제품여부"].map({"Y": "우수제품", "N": "일반제품"}).fillna("미표기")
        g = tmp.groupby("_우수")["공급금액"].agg(["sum", "count"]).reset_index()
        g.columns = ["구분", "금액", "건수"]
        a["우수제품"] = g
        excl = float(g.loc[g["구분"] == "우수제품", "금액"].sum())
        a["우수제품금액"] = excl
        a["우수제품비율"] = excl / total_amt * 100 if total_amt else 0

    # 지역별
    if "수요기관지역" in df and "공급금액" in df:
        tmp = df.copy()
        tmp["_시도"] = tmp["수요기관지역"].astype(str).str.split().str[0]
        g = tmp.groupby("_시도")["공급금액"].agg(["sum", "count"]).sort_values("sum", ascending=False)
        a["지역별"] = g.head(10).reset_index().rename(columns={"_시도": "지역"})

    # 계약방식별
    if "계약방법" in df and "공급금액" in df:
        g = df.groupby("계약방법")["공급금액"].agg(["sum", "count"]).sort_values("sum", ascending=False)
        a["계약방식"] = g.reset_index()

    # 기간 추이 (월별)
    if "계약일자_dt" in df and "공급금액" in df:
        base = df.dropna(subset=["계약일자_dt"]).set_index("계약일자_dt")["공급금액"]
        try:
            ts = base.resample("ME").sum()
        except ValueError:
            ts = base.resample("M").sum()
        a["월별추이"] = ts

    # 자동 인사이트 문장
    a["인사이트"] = _make_insights(a)
    return a


def _make_insights(a):
    """분석 데이터로 자동 총평 문장 생성"""
    lines = []
    term = a.get("검색어", "해당 품목")
    lines.append(
        f"조회 기간({a['기간']}) 동안 '{term}' 품목의 총 조달 규모는 "
        f"{_eok(a['총공급금액'])}이며, 총 {a['계약건수']:,}건의 계약이 체결되었습니다. "
        f"참여 업체는 {a['업체수']:,}개, 구매 수요기관은 {a['수요기관수']:,}개입니다."
    )

    if "상위업체" in a and len(a["상위업체"]) > 0:
        top = a["상위업체"].iloc[0]
        share = top["점유율"]
        msg = (f"공급 실적 1위 업체는 '{top['업체명']}'으로 "
               f"{_eok(top['sum'])}({share:.1f}%)를 기록했습니다.")
        if len(a["상위업체"]) >= 3:
            top3 = a["상위업체"].head(3)["점유율"].sum()
            msg += f" 상위 3개 업체가 전체의 {top3:.1f}%를 차지하고 있어,"
            if top3 >= 60:
                msg += " 시장이 소수 업체에 집중된 편입니다."
            elif top3 >= 35:
                msg += " 상위 업체 중심의 경쟁 구도를 보입니다."
            else:
                msg += " 비교적 다수 업체가 고르게 참여하고 있습니다."
        lines.append(msg)

    if "상위기관" in a and len(a["상위기관"]) > 0:
        top = a["상위기관"].iloc[0]
        lines.append(
            f"최대 수요기관은 '{top['수요기관']}'으로 {_eok(top['sum'])}"
            f"({top['점유율']:.1f}%)를 구매했습니다."
        )

    if "우수제품비율" in a:
        r = a["우수제품비율"]
        if r >= 50:
            tone = "우수제품이 시장의 과반을 차지하고 있습니다."
        elif r >= 20:
            tone = "우수제품이 상당 부분을 차지하고 있습니다."
        elif r > 0:
            tone = "우수제품 비중은 아직 낮은 편입니다."
        else:
            tone = "우수제품 계약은 확인되지 않았습니다."
        lines.append(f"금액 기준 우수제품 비중은 {r:.1f}%로, {tone}")

    if "지역별" in a and len(a["지역별"]) > 0:
        top = a["지역별"].iloc[0]
        lines.append(
            f"지역별로는 '{top['지역']}'의 수요가 {_eok(top['sum'])}로 가장 큽니다."
        )

    return lines


# ──────────────────────────────────────────────
# Word 문서 생성
# ──────────────────────────────────────────────
def build_docx(a, chart_images=None):
    """
    분석 데이터(a)로 Word 문서를 만들어 BytesIO로 반환.
    chart_images: [(제목, png_bytes), ...] 형태의 차트 이미지 리스트 (선택)
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # 제목
    title = doc.add_heading(level=0)
    run = title.add_run(f"조달 시장 분석 보고서")
    run.font.size = Pt(22)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"품목: {a['검색어']}    |    기간: {a['기간']}")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    gr = gen.add_run(f"작성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    gr.font.size = Pt(9)
    gr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 1. 핵심 요약
    doc.add_heading("1. 핵심 요약", level=1)
    summary_tbl = doc.add_table(rows=0, cols=2)
    summary_tbl.style = "Light Grid Accent 1"
    rows = [
        ("총 공급금액", _won(a["총공급금액"])),
        ("계약 건수", f"{a['계약건수']:,}건"),
        ("참여 업체 수", f"{a['업체수']:,}개"),
        ("수요기관 수", f"{a['수요기관수']:,}개"),
        ("평균 계약금액", _won(a["평균계약금액"])),
    ]
    for k, v in rows:
        cells = summary_tbl.add_row().cells
        cells[0].text = k
        cells[1].text = v

    # 2. 총평 / 인사이트
    doc.add_heading("2. 분석 총평", level=1)
    for line in a.get("인사이트", []):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(6)

    # 차트 이미지 (있으면)
    if chart_images:
        doc.add_heading("3. 주요 차트", level=1)
        for cap, img_bytes in chart_images:
            if img_bytes:
                try:
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.0))
                    cp = doc.add_paragraph(cap)
                    cp = cp.runs[0] if cp.runs else cp.add_run(cap)
                    cp.font.size = Pt(9)
                    cp.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                except Exception:
                    pass

    sec = 4 if chart_images else 3

    def add_table_section(title, dfo, col_specs):
        nonlocal sec
        if dfo is None or len(dfo) == 0:
            return
        doc.add_heading(f"{sec}. {title}", level=1)
        sec += 1
        t = doc.add_table(rows=1, cols=len(col_specs))
        t.style = "Light List Accent 1"
        hdr = t.rows[0].cells
        for i, (label, _, _) in enumerate(col_specs):
            hdr[i].text = label
        for _, row in dfo.iterrows():
            cells = t.add_row().cells
            for i, (_, key, fmt) in enumerate(col_specs):
                val = row[key]
                cells[i].text = fmt(val) if fmt else str(val)

    # 상위 업체
    if "상위업체" in a:
        add_table_section(
            "상위 공급업체", a["상위업체"],
            [("업체명", "업체명", None),
             ("공급금액", "sum", _won),
             ("계약건수", "count", lambda x: f"{int(x):,}건"),
             ("점유율", "점유율", lambda x: f"{x:.1f}%")],
        )

    # 상위 수요기관
    if "상위기관" in a:
        add_table_section(
            "상위 수요기관", a["상위기관"],
            [("수요기관", "수요기관", None),
             ("구매액", "sum", _won),
             ("계약건수", "count", lambda x: f"{int(x):,}건"),
             ("점유율", "점유율", lambda x: f"{x:.1f}%")],
        )

    # 우수제품
    if "우수제품" in a:
        add_table_section(
            "우수제품 비중", a["우수제품"],
            [("구분", "구분", None),
             ("금액", "금액", _won),
             ("건수", "건수", lambda x: f"{int(x):,}건")],
        )

    # 지역별
    if "지역별" in a:
        add_table_section(
            "지역별 분포", a["지역별"],
            [("지역", "지역", None),
             ("공급금액", "sum", _won),
             ("계약건수", "count", lambda x: f"{int(x):,}건")],
        )

    # 계약방식별
    if "계약방식" in a:
        add_table_section(
            "계약 방식별 분석", a["계약방식"],
            [("계약방법", "계약방법", None),
             ("공급금액", "sum", _won),
             ("계약건수", "count", lambda x: f"{int(x):,}건")],
        )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf